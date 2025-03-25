"""
Preview service for generating HTML and PDF previews of reports
"""

# Standard library imports
import os
import platform
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict

# Third-party imports
import mammoth
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import UUID4

# Local imports
from services.docx_service import docx_service
from utils.error_handler import handle_exception, logger
from utils.file_processor import FileProcessor
from utils.file_utils import safe_path_join

# Conditionally import Windows-specific modules
IS_WINDOWS = platform.system() == "Windows"
if IS_WINDOWS:
    try:
        import docx2pdf
        import pythoncom
    except ImportError:
        logger.warning(
            "pythoncom or docx2pdf not available, PDF conversion will be limited"
        )
else:
    logger.info("Running on non-Windows platform, using alternative PDF conversion")

class PreviewService:
    def __init__(self):
        self.preview_dir = Path("previews")
        self.preview_dir.mkdir(exist_ok=True)

        # Enhanced style map for better document structure
        self.style_map = """
        p[style-name='Title'] => h1.document-title:fresh
        p[style-name='Heading 1'] => h2.section-title:fresh
        p[style-name='Heading 2'] => h3.subsection-title:fresh
        p[style-name='Heading 3'] => h4.subsection-title:fresh
        p[style-name='Normal'] => p:fresh
        r[style-name='Strong'] => strong
        r[style-name='Emphasis'] => em
        table => table.document-table
        tr => tr
        td => td
        p[style-name='List Paragraph'] => li
        """

    async def generate_preview(self, report_id: UUID4) -> str:
        """
        Generate an enhanced HTML preview of a DOCX report.

        Args:
            report_id: UUID of the report to preview

        Returns:
            Path to the HTML preview file
        """
        try:
            # Get paths
            docx_path = docx_service.get_report_path(report_id)
            preview_path = docx_service.get_preview_path(report_id)

            # Convert DOCX to HTML with enhanced options
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=self.style_map,
                    convert_image=mammoth.images.img_element(
                        lambda image: {
                            "src": f"data:{image.content_type};base64,{image.base64_bytes.decode('utf-8')}",
                            "class": "document-image",
                        }
                    ),
                )

            # Add enhanced styles and responsive design
            html_content = f"""
            <!DOCTYPE html>
            <html lang="it">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Anteprima Report</title>
                <style>
                    :root {{
                        --primary-color: #2c3e50;
                        --secondary-color: #34495e;
                        --background-color: #f8f9fa;
                        --text-color: #2c3e50;
                        --border-color: #dee2e6;
                    }}
                    
                    body {{
                        font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
                        line-height: 1.6;
                        max-width: 1000px;
                        margin: 0 auto;
                        padding: 2rem;
                        background-color: var(--background-color);
                        color: var(--text-color);
                    }}
                    
                    .preview-container {{
                        background-color: white;
                        padding: 3rem;
                        border-radius: 8px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }}
                    
                    .document-title {{
                        color: var(--primary-color);
                        font-size: 2.5rem;
                        margin-bottom: 2rem;
                        text-align: center;
                        border-bottom: 2px solid var(--border-color);
                        padding-bottom: 1rem;
                    }}
                    
                    .section-title {{
                        color: var(--primary-color);
                        font-size: 1.8rem;
                        margin-top: 2.5rem;
                        margin-bottom: 1.5rem;
                        border-bottom: 1px solid var(--border-color);
                        padding-bottom: 0.5rem;
                    }}
                    
                    .subsection-title {{
                        color: var(--secondary-color);
                        font-size: 1.4rem;
                        margin-top: 2rem;
                        margin-bottom: 1rem;
                    }}
                    
                    p {{
                        margin-bottom: 1rem;
                        text-align: justify;
                    }}
                    
                    .document-table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 1.5rem 0;
                        background-color: white;
                    }}
                    
                    .document-table th,
                    .document-table td {{
                        border: 1px solid var(--border-color);
                        padding: 12px;
                        text-align: left;
                    }}
                    
                    .document-table th {{
                        background-color: var(--background-color);
                        font-weight: 600;
                    }}
                    
                    .document-table tr:nth-child(even) {{
                        background-color: var(--background-color);
                    }}
                    
                    .document-image {{
                        max-width: 100%;
                        height: auto;
                        margin: 1.5rem 0;
                        border-radius: 4px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }}
                    
                    @media print {{
                        body {{
                            background-color: white;
                            padding: 0;
                        }}
                        
                        .preview-container {{
                            box-shadow: none;
                            padding: 0;
                        }}
                    }}
                    
                    @media (max-width: 768px) {{
                        body {{
                            padding: 1rem;
                        }}
                        
                        .preview-container {{
                            padding: 1.5rem;
                        }}
                        
                        .document-title {{
                            font-size: 2rem;
                        }}
                        
                        .section-title {{
                            font-size: 1.5rem;
                        }}
                        
                        .subsection-title {{
                            font-size: 1.2rem;
                        }}
                    }}
                </style>
            </head>
            <body>
                <div class="preview-container">
                    {result.value}
                </div>
                <script>
                    // Add table responsiveness
                    document.querySelectorAll('table').forEach(table => {{
                        const wrapper = document.createElement('div');
                        wrapper.style.overflowX = 'auto';
                        table.parentNode.insertBefore(wrapper, table);
                        wrapper.appendChild(table);
                    }});
                </script>
            </body>
            </html>
            """

            # Save the HTML preview
            preview_path.write_text(html_content, encoding="utf-8")

            # Log any warnings
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Preview warning for {report_id}: {message}")

            logger.info(f"Generated enhanced preview for report {report_id}")
            return str(preview_path)

        except Exception as e:
            handle_exception(e, "Preview generation")
            raise

    def cleanup_old_previews(self, max_age_hours: int = 24):
        """
        Clean up preview files older than the specified age.

        Args:
            max_age_hours: Maximum age in hours before files are deleted
        """
        try:
            import time

            current_time = time.time()
            files_cleaned = 0

            # Make sure the preview directory exists
            if not self.preview_dir.exists():
                logger.warning("Preview directory does not exist, nothing to clean up")
                return

            for file_name in os.listdir(self.preview_dir):
                try:
                    # Use safe path joining to prevent directory traversal
                    preview_file = safe_path_join(self.preview_dir, file_name)

                    # Only process HTML files
                    if preview_file.suffix.lower() != ".html":
                        continue

                    # Check file age
                    file_age_hours = (
                        current_time - preview_file.stat().st_mtime
                    ) / 3600
                    if file_age_hours > max_age_hours:
                        preview_file.unlink()
                        files_cleaned += 1
                        logger.info(
                            f"Cleaned up old preview: {preview_file.name} (Age: {file_age_hours:.1f} hours)"
                        )
                except ValueError as e:
                    logger.warning(f"Skipping invalid preview file path: {e}")
                except Exception as e:
                    logger.error(
                        f"Error cleaning up preview file {file_name}: {str(e)}"
                    )

            if files_cleaned > 0:
                logger.info(
                    f"Preview cleanup complete. Removed {files_cleaned} old preview files."
                )

        except Exception as e:
            logger.error(f"Error cleaning up previews: {str(e)}")


def generate_docx_preview(
    content: str, include_header: bool = True, include_footer: bool = True
) -> Dict[str, Any]:
    """
    Generate a DOCX preview file for the given content.

    Args:
        content: Report content
        include_header: Whether to include header
        include_footer: Whether to include footer

    Returns:
        Dictionary with file paths and other preview info
    """
    try:
        # Create a temporary directory for our files
        temp_dir = tempfile.mkdtemp(prefix="report_preview_")
        docx_path = os.path.join(temp_dir, "preview.docx")
        pdf_path = os.path.join(temp_dir, "preview.pdf")

        # Create a new Document
        doc = Document()

        # Set up document style
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Add header if requested
        if include_header:
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "Insurance Report Preview"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add footer if requested
        if include_footer:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "PREVIEW - NOT FINAL DOCUMENT"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content to document
        for line in content.split("\n"):
            doc.add_paragraph(line)

        # Save the document
        doc.save(docx_path)
        logger.info(f"Generated DOCX preview at {docx_path}")

        # Generate PDF from DOCX
        pdf_data = None
        try:
            if IS_WINDOWS:
                # Windows-specific conversion using docx2pdf
                pythoncom.CoInitialize()  # Initialize COM for thread safety
                logger.info(
                    f"Converting DOCX to PDF using docx2pdf: {docx_path} -> {pdf_path}"
                )
                docx2pdf.convert(docx_path, pdf_path)
                logger.info(f"Generated PDF preview at {pdf_path}")

                # Get base64 encoded PDF data
                pdf_data = FileProcessor.get_file_as_base64(pdf_path)
            else:
                # Non-Windows conversion alternatives
                logger.info(
                    "Attempting alternative PDF conversion for non-Windows platform"
                )

                # Try LibreOffice if available
                try:
                    logger.info("Attempting conversion with LibreOffice")
                    cmd = [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        temp_dir,
                        docx_path,
                    ]
                    result = subprocess.run(
                        cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60
                    )

                    if result.returncode == 0:
                        # LibreOffice may use a different naming convention
                        pdf_name = os.path.basename(docx_path).replace(".docx", ".pdf")
                        pdf_path = os.path.join(temp_dir, pdf_name)
                        logger.info(
                            f"Generated PDF preview using LibreOffice at {pdf_path}"
                        )
                        pdf_data = FileProcessor.get_file_as_base64(pdf_path)
                    else:
                        logger.warning(
                            f"LibreOffice conversion failed: {result.stderr.decode('utf-8')}"
                        )
                        raise Exception("LibreOffice conversion failed")
                except Exception as libreoffice_error:
                    logger.warning(
                        f"LibreOffice conversion failed: {str(libreoffice_error)}"
                    )

                    # Try using unoconv as an alternative
                    try:
                        logger.info("Attempting conversion with unoconv")
                        cmd = ["unoconv", "-f", "pdf", "-o", pdf_path, docx_path]
                        result = subprocess.run(
                            cmd,
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            timeout=60,
                        )

                        if result.returncode == 0:
                            logger.info(
                                f"Generated PDF preview using unoconv at {pdf_path}"
                            )
                            pdf_data = FileProcessor.get_file_as_base64(pdf_path)
                        else:
                            logger.warning(
                                f"unoconv conversion failed: {result.stderr.decode('utf-8')}"
                            )
                            raise Exception("unoconv conversion failed")
                    except Exception as unoconv_error:
                        logger.warning(
                            f"unoconv conversion failed: {str(unoconv_error)}"
                        )

                        # If all else fails, indicate PDF generation isn't available
                        logger.info("PDF conversion not available on this platform")
                        pdf_path = None

        except Exception as pdf_error:
            logger.error(f"Error generating PDF preview: {str(pdf_error)}")
            pdf_path = None
            pdf_data = None

        # Get base64 encoded DOCX data
        docx_data = FileProcessor.get_file_as_base64(docx_path)

        return {
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "docx_data": docx_data,
            "pdf_data": pdf_data,
            "temp_dir": temp_dir,
        }

    except Exception as e:
        logger.error(f"Error generating document preview: {str(e)}")
        return {
            "docx_path": None,
            "pdf_path": None,
            "docx_data": None,
            "pdf_data": None,
            "temp_dir": None,
            "error": str(e),
        }


# Create singleton instance
preview_service = PreviewService()
