import asyncio
import datetime
import hashlib
import os
import re
import shutil
import tempfile
import time
import traceback
import uuid as uuid_lib
import logging
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import the resource manager 
from utils.resource_manager import resource_manager

# Setup logger directly instead of importing from error_handler
logger = logging.getLogger(__name__)

# Import metrics collector
from utils.metrics import MetricsCollector

# Import our service
from .docx_service import docx_service

# Simple local exception handler to avoid circular imports
def local_handle_exception(e: Exception, operation: str) -> None:
    """
    Local exception handler to avoid circular imports with utils.error_handler
    
    Args:
        e: The exception to handle
        operation: Description of the operation that failed
    """
    logger.error(f"Error in {operation}: {str(e)}")
    logger.error(f"Stack trace: {traceback.format_exc()}")


def parse_markdown(doc: Document, markdown_text: str) -> None:
    """
    Parse markdown text and add it to the document with proper formatting.

    Args:
        doc: Document to add content to
        markdown_text: Markdown text to parse
    """
    try:
        # Split text into paragraphs
        paragraphs = markdown_text.split("\n\n")

        for paragraph in paragraphs:
            if not paragraph.strip():
                continue

            # Handle headings
            if paragraph.startswith("#"):
                level = len(re.match("^#+", paragraph).group())
                text = paragraph.lstrip("#").strip()
                doc.add_heading(text, level=level)
                continue

            # Handle lists
            if paragraph.strip().startswith("- "):
                items = paragraph.strip().split("\n")
                for item in items:
                    if item.strip().startswith("- "):
                        p = doc.add_paragraph()
                        p.style = "List Bullet"
                        p.add_run(item.strip("- "))
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            p.add_run(paragraph.strip())

    except Exception as e:
        local_handle_exception(e, "Markdown parsing")
        raise


def process_inline_formatting(paragraph, text):
    """
    Process inline markdown formatting (bold, italic, etc.)

    Args:
        paragraph: The paragraph object to add text to
        text: Text with inline markdown formatting
    """
    # Process bold and italic formatting
    parts = []
    current_pos = 0

    # Find all bold and italic patterns
    # Bold: **text**
    # Italic: *text*
    # Bold+Italic: ***text***
    pattern = r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)"

    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > current_pos:
            parts.append((text[current_pos : match.start()], False, False))

        # Process the match
        match_text = match.group(0)
        if match_text.startswith("***") and match_text.endswith("***"):
            # Bold and italic
            content = match_text[3:-3]
            parts.append((content, True, True))
        elif match_text.startswith("**") and match_text.endswith("**"):
            # Bold
            content = match_text[2:-2]
            parts.append((content, True, False))
        elif match_text.startswith("*") and match_text.endswith("*"):
            # Italic
            content = match_text[1:-1]
            parts.append((content, False, True))

        current_pos = match.end()

    # Add any remaining text
    if current_pos < len(text):
        parts.append((text[current_pos:], False, False))

    # Add all parts to the paragraph with appropriate formatting
    for part_text, bold, italic in parts:
        run = paragraph.add_run(part_text)
        run.bold = bold
        run.italic = italic


def replace_template_variables(doc: Document, variables: Dict[str, Any]) -> None:
    """
    Replace all {{ variable }} placeholders in the document with their values.

    Args:
        doc: Document to process
        variables: Dictionary of variable names and values
    """
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                for var_name, var_value in variables.items():
                    placeholder = f"{{{{ {var_name} }}}}"
                    if placeholder in text:
                        text = text.replace(placeholder, str(var_value))
                run.text = text

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            for var_name, var_value in variables.items():
                                placeholder = f"{{{{ {var_name} }}}}"
                                if placeholder in text:
                                    text = text.replace(placeholder, str(var_value))
                            run.text = text

    except Exception as e:
        local_handle_exception(e, "Template variable replacement")
        raise


def generate_docx(
    report_text: str,
    output_filename: str,
    reference_metadata: Optional[Dict[str, Any]] = None,
    template_variables: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Create a DOCX file using the template, replacing variables and appending report text.

    Args:
        report_text: Text content to append to the document
        output_filename: Name of the output file
        reference_metadata: Optional metadata from reference reports
        template_variables: Optional variables to replace in the template

    Returns:
        Path to the generated file
    """
    try:
        # Use docx_service to handle template and variable replacement
        if template_variables:
            report_id = docx_service.generate_report(template_variables)
            doc_path = docx_service.get_report_path(report_id)
        else:
            # Create a new document if no template variables
            doc = Document()

            # Add title
            title = doc.add_heading("Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add report text
            parse_markdown(doc, report_text)

            # Save the document
            doc_path = Path(output_filename)
            doc.save(str(doc_path))

        return str(doc_path)

    except Exception as e:
        local_handle_exception(e, "DOCX generation")
        raise


def format_report_as_docx(
    report_content: str,
    reference_metadata: Optional[Dict[str, Any]] = None,
    filename: Optional[str] = None,
    template_variables: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Format the report content as a DOCX document.

    Args:
        report_content: Content to include in the report
        reference_metadata: Optional metadata from reference reports
        filename: Optional name for the output file
        template_variables: Optional variables to replace in the template

    Returns:
        Path to the generated file
    """
    try:
        if not filename:
            # Generate a unique filename
            from uuid import uuid4

            filename = f"report_{uuid4()}.docx"

        # Generate the DOCX file
        doc_path = generate_docx(
            report_content, filename, reference_metadata, template_variables
        )

        return doc_path

    except Exception as e:
        local_handle_exception(e, "Report formatting")
        raise


# Constants
TEMPLATE_DIR = Path(__file__).parent.parent / "data" / "templates"
OUTPUT_DIR = Path(__file__).parent.parent / "generated_reports"
CACHE_DIR = Path(__file__).parent.parent / "data" / "cache"
MAX_CACHE_SIZE_MB = 100  # Maximum cache size in MB
QUALITY_THRESHOLD = 0.7  # Threshold for document quality score

# Ensure directories exist
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
CACHE_DIR.mkdir(parents=True, exist_ok=True)


class DocxFormatter:
    """
    Handles DOCX document generation with template support, quality checks, and performance tracking.
    """

    def __init__(self):
        self.templates = self._load_templates()
        self.metrics = {
            "total_documents": 0,
            "avg_generation_time": 0,
            "cache_hits": 0,
            "cache_misses": 0,
            "quality_check_failures": 0,
        }
        self.metrics_lock = Lock()
        self.document_cache = {}
        self.cache_lock = Lock()

        # Initialize metrics collector
        metrics_file = Path(__file__).parent.parent / "data" / "metrics.json"
        self.metrics_collector = MetricsCollector(metrics_file=metrics_file)

    def _load_templates(self) -> Dict[str, Path]:
        """Load available templates from the template directory"""
        templates = {}
        for template_file in TEMPLATE_DIR.glob("*.docx"):
            template_type = (
                template_file.stem
            )  # Use filename without extension as template type
            templates[template_type] = template_file

        # Add default template if none exist
        if not templates:
            default_template = TEMPLATE_DIR / "default.docx"
            if not default_template.exists():
                document = Document()
                document.save(default_template)
            templates["default"] = default_template

        logger.info(f"Loaded {len(templates)} document templates")
        return templates

    def _get_template(self, template_type: str = "default") -> Path:
        """Get the appropriate template file"""
        if template_type in self.templates:
            return self.templates[template_type]

        logger.warning(f"Template '{template_type}' not found, using default")
        return self.templates.get(
            "default", self.templates[list(self.templates.keys())[0]]
        )

    def _get_cache_key(self, content: str, template_type: str) -> str:
        """Generate a cache key based on content and template type"""
        hash_input = f"{content}:{template_type}:{datetime.date.today().isoformat()}"
        return hashlib.md5(hash_input.encode()).hexdigest()

    def _check_cache(self, cache_key: str) -> Optional[Path]:
        """Check if a document exists in cache"""
        with self.cache_lock:
            if cache_key in self.document_cache:
                cached_path = self.document_cache[cache_key]

                # Verify the file exists
                if Path(cached_path).exists():
                    with self.metrics_lock:
                        self.metrics["cache_hits"] += 1
                    # Update the access time
                    Path(cached_path).touch()
                    return cached_path

                # Remove from cache if file doesn't exist
                del self.document_cache[cache_key]

            with self.metrics_lock:
                self.metrics["cache_misses"] += 1

            return None

    async def _clean_cache(self) -> None:
        """Clean up old files from cache directory"""
        # Get cache size
        total_size = sum(
            f.stat().st_size for f in CACHE_DIR.glob("**/*") if f.is_file()
        )
        total_size_mb = total_size / (1024 * 1024)

        if total_size_mb < MAX_CACHE_SIZE_MB:
            return

        logger.info(
            f"Cache size ({total_size_mb:.2f} MB) exceeds limit ({MAX_CACHE_SIZE_MB} MB), cleaning up"
        )

        # List files by access time
        files = [
            (f, f.stat().st_atime) for f in CACHE_DIR.glob("*.docx") if f.is_file()
        ]
        files.sort(key=lambda x: x[1])  # Sort by access time, oldest first

        # Delete oldest files until under limit
        for file_path, _ in files:
            file_size = file_path.stat().st_size / (1024 * 1024)
            file_path.unlink()
            total_size_mb -= file_size
            logger.debug(f"Deleted cached file: {file_path}")

            if (
                total_size_mb < MAX_CACHE_SIZE_MB * 0.8
            ):  # Stop if we're under 80% of max
                break

    def _add_document_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add metadata properties to the document"""
        core_properties = doc.core_properties
        core_properties.title = metadata.get("title", "Insurance Report")
        core_properties.author = metadata.get("author", "Insurance Report Generator")
        core_properties.created = metadata.get("created", datetime.datetime.now())
        core_properties.modified = datetime.datetime.now()
        core_properties.category = metadata.get("category", "Insurance")
        core_properties.comments = metadata.get("comments", "Generated document")

        # Add custom properties
        if "report_id" in metadata:
            self._add_custom_property(doc, "ReportId", metadata["report_id"])
        if "version" in metadata:
            self._add_custom_property(doc, "Version", metadata["version"])

    def _add_custom_property(self, doc: Document, name: str, value: str) -> None:
        """Add a custom property to the document"""
        try:
            custom_props = doc._part.custom_properties_part
            if custom_props is None:
                doc._part.add_custom_properties_part()
                custom_props = doc._part.custom_properties_part

            custom_props.props.add_property(name, value)
        except AttributeError:
            logger.warning(
                f"Failed to add custom property {name}, using python-docx that doesn't support it"
            )

    def _validate_document_quality(
        self, content: str, document: Document
    ) -> Tuple[bool, Dict[str, Any]]:
        """Check document quality and return (passed, quality_metrics)"""
        quality_metrics = {"score": 0.0, "issues": [], "passed": False}

        # Check 1: Content length
        if len(content) < 100:
            quality_metrics["issues"].append("Content is too short")

        # Check 2: Document structure (must have at least a title and one paragraph)
        if len(document.paragraphs) < 2:
            quality_metrics["issues"].append("Document has insufficient structure")

        # Check 3: Verify that all sections have content
        empty_sections = 0
        section_count = 0

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                section_count += 1
                if section_count > 1:  # Check if previous section was empty
                    if empty_sections > 0:
                        quality_metrics["issues"].append(
                            f"Empty or minimal content in section {section_count-1}"
                        )
                empty_sections = 0
            elif paragraph.text.strip():
                empty_sections = 0
            else:
                empty_sections += 1

        # Check 4: Basic grammar check (simplified)
        grammar_issues = self._check_basic_grammar(content)
        if grammar_issues:
            quality_metrics["issues"].extend(
                grammar_issues[:3]
            )  # Limit to top 3 issues

        # Calculate final score
        base_score = 1.0
        issue_penalty = 0.1 * len(quality_metrics["issues"])
        quality_metrics["score"] = max(0.0, min(1.0, base_score - issue_penalty))
        quality_metrics["passed"] = quality_metrics["score"] >= QUALITY_THRESHOLD

        if not quality_metrics["passed"]:
            self.metrics["quality_check_failures"] += 1

        return quality_metrics["passed"], quality_metrics

    def _check_basic_grammar(self, content: str) -> List[str]:
        """Perform basic grammar checks on content"""
        issues = []

        # Check for repeated words
        repeated_words = re.findall(r"\b(\w+)\s+\1\b", content, re.IGNORECASE)
        if repeated_words:
            issues.append(f"Repeated words detected: {', '.join(set(repeated_words))}")

        # Check for very long sentences (>50 words)
        sentences = re.split(r"[.!?]+", content)
        for i, sentence in enumerate(sentences):
            words = sentence.split()
            if len(words) > 50:
                issues.append(f"Very long sentence detected: sentence #{i+1}")

        # Check for very short paragraphs
        paragraphs = content.split("\n\n")
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip() and len(paragraph.split()) < 3:
                issues.append(f"Very short paragraph detected: paragraph #{i+1}")

        return issues

    async def generate_document(
        self,
        content: str,
        output_path: Optional[str] = None,
        template_type: str = "default",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Generate a DOCX document from content using a template

        Args:
            content: Document text content
            output_path: Optional path to save the document, auto-generated if None
            template_type: Template to use for document generation
            metadata: Optional metadata to include in the document

        Returns:
            Dict with document details including path, quality info, and metrics
        """
        start_time = time.time()
        self.metrics["total_documents"] += 1

        if metadata is None:
            metadata = {}

        # Generate a unique filename if not provided
        if output_path is None:
            safe_title = re.sub(r"[^\w\-_\.]", "_", metadata.get("title", "report"))[
                :50
            ]
            report_id = metadata.get(
                "report_id", hashlib.md5(content[:100].encode()).hexdigest()[:8]
            )
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{report_id}_{timestamp}.docx"
            output_path = str(OUTPUT_DIR / filename)

        # Check cache for existing document
        cache_key = self._get_cache_key(content, template_type)
        cached_doc = self._check_cache(cache_key)

        if cached_doc:
            logger.info(f"Using cached document for key {cache_key}")
            # Copy cached document to output path
            shutil.copy(cached_doc, output_path)

            generation_time = time.time() - start_time
            return {
                "path": output_path,
                "url": f"/reports/{os.path.basename(output_path)}",
                "from_cache": True,
                "generation_time": generation_time,
                "quality": {
                    "score": 1.0,
                    "passed": True,
                },  # Assume cached documents passed quality check
            }

        # Begin document creation
        result = None
        temp_files = []

        try:
            # Create a document from template
            template_path = self._get_template(template_type)
            document = Document(template_path)

            # Add content with proper formatting
            self._format_document(document, content)

            # Add metadata
            timestamp = datetime.datetime.now()
            default_metadata = {
                "created": timestamp,
                "version": "1.0",
                "generator": "Insurance Report Generator",
            }

            # Merge default with provided metadata
            full_metadata = {**default_metadata, **metadata}
            self._add_document_metadata(document, full_metadata)

            # Perform quality checks
            passed_quality, quality_metrics = self._validate_document_quality(
                content, document
            )

            # Create atomic save (save to temp file first, then move)
            temp_file_id = str(uuid_lib.uuid4())
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                document.save(temp_file.name)
                temp_path = temp_file.name
                # Register the temp file for cleanup
                resource_manager.register_resource("temp_file", temp_file_id, temp_path)
                temp_files.append(temp_file_id)

            # Move the temp file to the final location
            os.replace(temp_path, output_path)

            # Save to cache if quality check passed
            if passed_quality:
                cache_path = CACHE_DIR / f"{cache_key}.docx"
                shutil.copy(output_path, cache_path)

            # Calculate metrics
            generation_time = time.time() - start_time
            avg_time = self.metrics["avg_generation_time"]
            doc_count = self.metrics["total_documents"]
            self.metrics["avg_generation_time"] = (
                avg_time * (doc_count - 1) + generation_time
            ) / doc_count

            # Schedule cache cleaning if this isn't a cached document
            asyncio.create_task(self._clean_cache())

            # Prepare result
            result = {
                "path": output_path,
                "url": f"/reports/{os.path.basename(output_path)}",
                "from_cache": False,
                "generation_time": generation_time,
                "quality": quality_metrics,
            }

            logger.info(
                f"Document generated in {generation_time:.2f}s with quality score {quality_metrics['score']:.2f}"
            )
            return result

        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            if result is None:
                result = {
                    "error": str(e),
                    "path": None,
                    "generation_time": time.time() - start_time,
                }
            return result
        finally:
            # Clean up any registered temp files
            for temp_id in temp_files:
                resource_manager.cleanup_resource("temp_file", temp_id)

    def _format_document(self, document: Document, content: str) -> None:
        """Format the document with the provided content"""
        # Clear existing content (keep only first paragraph if it looks like a title)
        if (
            document.paragraphs
            and "heading" in document.paragraphs[0].style.name.lower()
        ):
            for i in range(1, len(document.paragraphs)):
                p = document.paragraphs[i]
                p.clear()
        else:
            for p in document.paragraphs:
                p.clear()

        # Process the content - split by markdown-style headers
        sections = re.split(r"(^|\n)(#{1,6})\s+([^\n]+)", content, flags=re.MULTILINE)

        # If no sections found, add content directly
        if len(sections) <= 1:
            p = document.add_paragraph(content)
            p.style = "Normal"
            return

        # Process each section with headers
        current_text = sections[0].strip()
        if current_text:
            p = document.add_paragraph(current_text)
            p.style = "Normal"

        i = 1
        while i < len(sections):
            if sections[i] and sections[i + 1] and sections[i + 2]:
                # This is a header + content pattern
                level = len(sections[i + 1].strip())  # Number of # characters
                header_text = sections[i + 2].strip()

                # Add header
                h = document.add_paragraph(header_text)
                h.style = f"Heading {min(level, 6)}"

                # Add content if available
                if i + 3 < len(sections) and sections[i + 3]:
                    content_text = sections[i + 3].strip()
                    if content_text:
                        p = document.add_paragraph(content_text)
                        p.style = "Normal"

                i += 3
            else:
                i += 1

        # Add page numbers
        self._add_page_numbers(document)

    def _add_page_numbers(self, document: Document) -> None:
        """Add page numbers to the document footer"""
        for section in document.sections:
            footer = section.footer
            footer_para = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add page number field
            run = footer_para.add_run()
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            run._r.append(fldChar1)

            instrText = OxmlElement("w:instrText")
            instrText.text = " PAGE "
            run._r.append(instrText)

            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar2)

    def get_metrics(self) -> Dict[str, Any]:
        """Get generation metrics"""
        with self.metrics_lock:
            return self.metrics.copy()


# Create a singleton instance
docx_formatter = DocxFormatter()


async def generate_docx_async(
    content: str,
    output_path: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Generate a DOCX document from the given content

    Args:
        content: The text content for the document
        output_path: Optional path to save the document
        metadata: Optional metadata to include in the document

    Returns:
        Dictionary with path to generated document and other information
    """
    if metadata is None:
        metadata = {}

    result = await docx_formatter.generate_document(
        content=content,
        output_path=output_path,
        template_type=metadata.get("template_type", "default"),
        metadata=metadata,
    )

    return result


async def get_document_metrics() -> Dict[str, Any]:
    """Get document generation metrics"""
    return docx_formatter.get_metrics()
